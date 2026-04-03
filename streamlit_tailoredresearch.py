import json
from io import BytesIO
from html import escape

import streamlit as st
from docx import Document

from database import SessionLocal
from tailoredresearch import conduct_research


def _build_word_export(result: dict, note_name: str, research_focus: str, source: str, limit: int) -> bytes:
    doc = Document()
    doc.add_heading("Research Hub Export", 0)

    doc.add_heading("Query", level=1)
    doc.add_paragraph(f"Note name: {note_name}")
    doc.add_paragraph(f"Research focus: {research_focus}")
    doc.add_paragraph(f"Source: {source}")
    doc.add_paragraph(f"Limit: {limit}")
    doc.add_paragraph(f"Timestamp: {result.get('timestamp', 'N/A')}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Papers found: {result.get('papers_found', 0)}")
    doc.add_paragraph("Sources queried: " + ", ".join(result.get("sources_queried", [])))
    doc.add_paragraph(f"Primary reference: {result.get('primary_reference', 'N/A')}")

    papers = result.get("papers")
    doc.add_heading("Papers", level=1)
    if isinstance(papers, list) and papers:
        for idx, paper in enumerate(papers, start=1):
            doc.add_heading(f"{idx}. {paper.get('title', 'No Title')}", level=2)
            doc.add_paragraph(f"Source: {paper.get('source', 'unknown')}")
            doc.add_paragraph(f"Journal: {paper.get('journal', 'Unknown Journal')}")
            doc.add_paragraph(f"Year: {paper.get('year', 'Unknown Year')}")
            authors = paper.get("authors") or []
            doc.add_paragraph("Authors: " + (", ".join(authors) if authors else "N/A"))
            doc.add_paragraph(f"Link: {paper.get('link') or 'N/A'}")
            doc.add_paragraph("Summary:")
            doc.add_paragraph(str(paper.get("summary", "No Abstract available")))
    else:
        doc.add_paragraph("No relevant papers found for this topic.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

st.set_page_config(page_title="Research Hub", page_icon="🔬", layout="wide")

st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@500;700;800&display=swap');

        :root {
            --rh-bg-0: #050505;
            --rh-bg-1: #0d0d0d;
            --rh-card: #161616;
            --rh-card-soft: #202020;
            --rh-text: #f6f6f6;
            --rh-muted: #adadad;
            --rh-accent: #1ed760;
            --rh-accent-soft: #15b24f;
            --rh-pink: #ff4fd8;
            --rh-cyan: #2dd4ff;
            --rh-amber: #f7b500;
        }

        .stApp {
            background:
                radial-gradient(circle at 12% 8%, rgba(45, 212, 255, 0.22) 0%, rgba(45, 212, 255, 0) 34%),
                radial-gradient(circle at 85% 6%, rgba(255, 79, 216, 0.2) 0%, rgba(255, 79, 216, 0) 38%),
                radial-gradient(circle at 50% 82%, rgba(30, 215, 96, 0.17) 0%, rgba(30, 215, 96, 0) 32%),
                linear-gradient(180deg, var(--rh-bg-1) 0%, var(--rh-bg-0) 100%);
            color: var(--rh-text);
            font-family: 'Montserrat', sans-serif;
        }

        [data-testid="stHeader"] {
            background: transparent;
        }

        [data-testid="stSidebar"] {
            border-right: 1px solid #2a2a2a;
            background: linear-gradient(180deg, #111111 0%, #0a0a0a 100%);
        }

        div[data-testid="stForm"] {
            border-radius: 18px;
            border: 1px solid #2b2b2b;
            background: rgba(20, 20, 20, 0.84);
            backdrop-filter: blur(5px);
            padding: 1rem 1.2rem 0.4rem 1.2rem;
            margin-bottom: 1rem;
        }

        div[data-testid="stMetric"] {
            border: 1px solid #323232;
            border-radius: 16px;
            background: var(--rh-card);
            padding: 0.8rem;
        }

        .rh-hero {
            background: linear-gradient(120deg, #1db954 0%, #2dd4ff 42%, #ff4fd8 100%);
            border-radius: 22px;
            padding: 1.6rem;
            border: 1px solid #2f2f2f;
            box-shadow: 0 8px 28px rgba(0, 0, 0, 0.35);
            margin-bottom: 1rem;
        }

        .rh-hero h1 {
            margin: 0;
            font-size: 2rem;
            font-weight: 800;
            letter-spacing: 0.4px;
        }

        .rh-hero p {
            margin-top: 0.45rem;
            margin-bottom: 0;
            color: #f3fff7;
        }

        .rh-card {
            border: 1px solid #3a3a3a;
            background:
                linear-gradient(180deg, var(--rh-card-soft) 0%, var(--rh-card) 100%),
                linear-gradient(90deg, rgba(45, 212, 255, 0.12) 0%, rgba(255, 79, 216, 0.1) 100%);
            border-radius: 16px;
            padding: 1rem;
            margin-bottom: 0.8rem;
            box-shadow: 0 6px 18px rgba(0, 0, 0, 0.28);
        }

        .rh-title {
            font-weight: 700;
            font-size: 1.05rem;
            margin: 0;
        }

        .rh-meta {
            color: var(--rh-muted);
            font-size: 0.88rem;
            margin-top: 0.35rem;
            margin-bottom: 0.55rem;
        }

        .rh-badge {
            display: inline-block;
            border-radius: 999px;
            padding: 0.2rem 0.55rem;
            font-size: 0.76rem;
            font-weight: 700;
            margin-right: 0.45rem;
            color: #0a0a0a;
            background: var(--rh-accent);
        }

        .rh-badge.europe_pmc { background: var(--rh-cyan); }
        .rh-badge.openalex { background: var(--rh-pink); }
        .rh-badge.pubmed { background: var(--rh-amber); }

        .rh-link a {
            color: var(--rh-accent) !important;
            text-decoration: none;
            font-weight: 700;
        }

        .rh-link a:hover {
            color: #8af8af !important;
        }

        .stButton > button,
        [data-testid="stFormSubmitButton"] > button {
            border-radius: 999px;
            border: none;
            font-weight: 800;
            background: linear-gradient(90deg, var(--rh-accent) 0%, var(--rh-accent-soft) 100%);
            color: #0b0b0b;
            transition: transform 120ms ease, box-shadow 120ms ease;
        }

        .stButton > button:hover,
        [data-testid="stFormSubmitButton"] > button:hover {
            transform: translateY(-1px);
            box-shadow: 0 8px 22px rgba(30, 215, 96, 0.32);
        }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    """
    <div class="rh-hero">
      <h1>Research Hub</h1>
      <p>Discovery flow for PubMed, Europe PMC, and OpenAlex.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

left_col, right_col = st.columns([2.0, 1.0], gap="large")

with left_col:
    with st.form("research_form"):
        note_name = st.text_input("Note name", placeholder="e.g., mRNA delivery optimization")
        research_focus = st.text_area(
            "Research focus",
            placeholder="Describe the topic or paste keywords.",
            height=145,
        )
        source = st.selectbox(
            "Data source",
            options=["both", "europe_pmc", "openalex", "pubmed"],
            index=0,
            help="Use both to query Europe PMC and OpenAlex in one run.",
        )
        limit = st.slider("Max papers", min_value=1, max_value=20, value=6)
        submitted = st.form_submit_button("Run research")

with right_col:
    st.markdown("### Sources")
    st.markdown(
        """
        <div class="rh-card">
          <span class="rh-badge pubmed">pubmed</span>
          <span class="rh-badge europe_pmc">europe_pmc</span>
          <span class="rh-badge openalex">openalex</span>
          <p class="rh-meta" style="margin-top:0.7rem;">Use the selector to query one source or combine multiple in one run.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

if submitted:
    if not note_name.strip() or not research_focus.strip():
        st.error("Please provide both a note name and a research focus.")
    else:
        db = SessionLocal()
        try:
            with st.spinner("Querying selected source(s) and saving note..."):
                result = conduct_research(
                    note_name=note_name.strip(),
                    research_focus=research_focus.strip(),
                    limit=limit,
                    source=source,
                    db=db,
                )

            st.success(f"Saved note: {result['note_saved_as']}")

            m1, m2 = st.columns(2)
            with m1:
                st.metric("Papers found", result.get("papers_found", 0))
            with m2:
                st.metric("Sources", len(result.get("sources_queried", [])))

            st.write("Sources queried: " + ", ".join(result.get("sources_queried", [])))

            export_bytes = _build_word_export(
                result=result,
                note_name=note_name.strip(),
                research_focus=research_focus.strip(),
                source=source,
                limit=limit,
            )
            export_name = f"{note_name.strip().replace(' ', '_') or 'research_export'}.docx"
            st.download_button(
                label="Export as Word (.docx)",
                data=export_bytes,
                file_name=export_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            primary_reference = result.get("primary_reference")
            if primary_reference and primary_reference.startswith("http"):
                st.markdown(f"Primary reference: [{primary_reference}]({primary_reference})")
            else:
                st.write(f"Primary reference: {primary_reference}")

            papers = result.get("papers")
            if isinstance(papers, list) and papers:
                st.subheader("Top papers")
                for paper in papers:
                    title = escape(paper.get("title", "No Title"))
                    source_name = escape(str(paper.get("source", "unknown")))
                    journal = escape(str(paper.get("journal", "Unknown Journal")))
                    year = escape(str(paper.get("year", "Unknown Year")))
                    authors = paper.get("authors") or []
                    authors_str = escape(", ".join(authors[:7]) + ("..." if len(authors) > 7 else "")) if authors else "N/A"
                    summary = escape(str(paper.get("summary", "No Abstract available")))
                    link = paper.get("link")

                    st.markdown(
                        f"""
                        <div class="rh-card">
                            <span class="rh-badge {source_name}">{source_name}</span>
                            <p class="rh-title">{title}</p>
                            <p class="rh-meta">{journal} | {year} | {authors_str}</p>
                            <p>{summary}</p>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )
                    if link:
                        safe_link = escape(str(link), quote=True)
                        st.markdown(
                            f"<p class='rh-link'><a href='{safe_link}' target='_blank'>Open paper</a></p>",
                            unsafe_allow_html=True,
                        )
            else:
                st.info("No relevant papers found for this topic.")

            with st.expander("Raw output JSON"):
                st.code(json.dumps(result, indent=2, ensure_ascii=False), language="json")
        finally:
            db.close()
